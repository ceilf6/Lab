"""
Word Document MCP Server

A simple MCP server for Word document operations including:
- Create, read, update, delete documents
- Add tables and images
- Format text and search/replace
- Merge documents
"""

from mcp.server.fastmcp import FastMCP
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import json
from pathlib import Path
from datetime import datetime
from typing import Optional, List

# Initialize FastMCP server
mcp = FastMCP("Word Document MCP Server")

# Default word directory
WORD_DIR = Path("word")
WORD_DIR.mkdir(exist_ok=True)


# ==================== Helper Functions ====================

def get_file_path(filename: str) -> Path:
    """Get full path for a file in word directory."""
    path = Path(filename)
    if not path.is_absolute():
        path = WORD_DIR / filename
    if not str(path).endswith('.docx'):
        path = Path(str(path) + '.docx')
    return path


# ==================== Tools ====================

@mcp.tool()
def create_document(
    filename: Optional[str] = None,
    title: Optional[str] = None,
    content: Optional[str] = None
) -> dict:
    """
    Create a new Word document.
    
    Args:
        filename: File name (optional, auto-generated if not provided)
        title: Document title (optional)
        content: Initial content (optional)
    
    Returns:
        Result dictionary with file path and status
    """
    try:
        # Generate filename if not provided
        if not filename:
            filename = f"document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        file_path = get_file_path(filename)
        
        # Create document
        doc = Document()
        
        if title:
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if content:
            for line in content.split('\n'):
                if line.strip():
                    doc.add_paragraph(line.strip())
                else:
                    doc.add_paragraph()
        
        doc.save(str(file_path))
        
        return {
            "success": True,
            "message": "Document created successfully",
            "file_path": str(file_path),
            "file_size": file_path.stat().st_size
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@mcp.tool()
def read_document(filename: str) -> dict:
    """
    Read content from a Word document.
    
    Args:
        filename: File name to read
    
    Returns:
        Document content and metadata
    """
    try:
        file_path = get_file_path(filename)
        
        if not file_path.exists():
            return {"success": False, "error": f"File not found: {filename}"}
        
        doc = Document(str(file_path))
        
        # Extract paragraphs
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        # Extract tables
        tables = []
        for table in doc.tables:
            table_data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            tables.append(table_data)
        
        return {
            "success": True,
            "file_path": str(file_path),
            "paragraphs": paragraphs,
            "paragraph_count": len(paragraphs),
            "tables": tables,
            "table_count": len(tables),
            "full_text": "\n".join(paragraphs)
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@mcp.tool()
def update_document(
    filename: str,
    action: str,
    content: Optional[str] = None,
    paragraph_index: Optional[int] = None
) -> dict:
    """
    Update a Word document.
    
    Args:
        filename: File name to update
        action: Action type - "append", "insert", "replace", "add_heading"
        content: Content to add/insert/replace
        paragraph_index: Paragraph index for insert/replace (0-based)
    
    Returns:
        Result dictionary
    """
    try:
        file_path = get_file_path(filename)
        
        if not file_path.exists():
            return {"success": False, "error": f"File not found: {filename}"}
        
        doc = Document(str(file_path))
        
        if action == "append":
            for line in content.split('\n') if content else []:
                if line.strip():
                    doc.add_paragraph(line.strip())
        
        elif action == "add_heading":
            if content:
                doc.add_heading(content, level=2)
        
        elif action == "insert" and paragraph_index is not None:
            if content and paragraph_index < len(doc.paragraphs):
                doc.paragraphs[paragraph_index].insert_paragraph_before(content)
        
        elif action == "replace" and paragraph_index is not None:
            if paragraph_index < len(doc.paragraphs):
                para = doc.paragraphs[paragraph_index]
                para.clear()
                para.add_run(content or "")
        
        else:
            return {"success": False, "error": f"Invalid action or missing parameters"}
        
        doc.save(str(file_path))
        
        return {
            "success": True,
            "message": "Document updated successfully",
            "action": action
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@mcp.tool()
def delete_document(filename: str) -> dict:
    """
    Delete a Word document.
    
    Args:
        filename: File name to delete
    
    Returns:
        Result dictionary
    """
    try:
        file_path = get_file_path(filename)
        
        if not file_path.exists():
            return {"success": False, "error": f"File not found: {filename}"}
        
        file_path.unlink()
        
        return {
            "success": True,
            "message": "Document deleted successfully"
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@mcp.tool()
def list_documents() -> dict:
    """
    List all Word documents in the word directory.
    
    Returns:
        List of documents with metadata
    """
    try:
        docs = []
        for file in WORD_DIR.glob("*.docx"):
            stat = file.stat()
            docs.append({
                "name": file.name,
                "path": str(file),
                "size": stat.st_size,
                "modified": datetime.fromtimestamp(stat.st_mtime).isoformat()
            })
        
        return {
            "success": True,
            "count": len(docs),
            "documents": docs
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@mcp.tool()
def add_table(
    filename: str,
    table_data: List[List[str]],
    title: Optional[str] = None
) -> dict:
    """
    Add a table to a document.
    
    Args:
        filename: File name
        table_data: 2D list of table data
        title: Optional table title
    
    Returns:
        Result dictionary
    """
    try:
        file_path = get_file_path(filename)
        
        if not file_path.exists():
            return {"success": False, "error": f"File not found: {filename}"}
        
        doc = Document(str(file_path))
        
        if title:
            doc.add_heading(title, level=2)
        
        if table_data:
            rows = len(table_data)
            cols = max(len(row) for row in table_data)
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Light Grid Accent 1'
            
            for i, row_data in enumerate(table_data):
                for j, cell_data in enumerate(row_data):
                    if j < cols:
                        table.rows[i].cells[j].text = str(cell_data)
        
        doc.save(str(file_path))
        
        return {
            "success": True,
            "message": "Table added successfully"
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@mcp.tool()
def insert_image(
    filename: str,
    image_path: str,
    width: Optional[float] = None
) -> dict:
    """
    Insert an image into a document.
    
    Args:
        filename: Document file name
        image_path: Path to image file
        width: Image width in inches (optional)
    
    Returns:
        Result dictionary
    """
    try:
        file_path = get_file_path(filename)
        
        if not file_path.exists():
            return {"success": False, "error": f"File not found: {filename}"}
        
        if not Path(image_path).exists():
            return {"success": False, "error": f"Image not found: {image_path}"}
        
        doc = Document(str(file_path))
        
        if width:
            doc.add_picture(image_path, width=Inches(width))
        else:
            doc.add_picture(image_path)
        
        doc.save(str(file_path))
        
        return {
            "success": True,
            "message": "Image inserted successfully"
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@mcp.tool()
def format_text(
    filename: str,
    paragraph_index: int,
    bold: bool = False,
    italic: bool = False,
    font_size: Optional[int] = None
) -> dict:
    """
    Format text in a paragraph.
    
    Args:
        filename: Document file name
        paragraph_index: Paragraph index (0-based)
        bold: Make text bold
        italic: Make text italic
        font_size: Font size in points
    
    Returns:
        Result dictionary
    """
    try:
        file_path = get_file_path(filename)
        
        if not file_path.exists():
            return {"success": False, "error": f"File not found: {filename}"}
        
        doc = Document(str(file_path))
        
        if paragraph_index >= len(doc.paragraphs):
            return {"success": False, "error": "Paragraph index out of range"}
        
        para = doc.paragraphs[paragraph_index]
        
        for run in para.runs:
            if bold:
                run.font.bold = True
            if italic:
                run.font.italic = True
            if font_size:
                run.font.size = Pt(font_size)
        
        doc.save(str(file_path))
        
        return {
            "success": True,
            "message": "Text formatted successfully"
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@mcp.tool()
def search_replace(
    filename: str,
    search_text: str,
    replace_text: str
) -> dict:
    """
    Search and replace text in a document.
    
    Args:
        filename: Document file name
        search_text: Text to search for
        replace_text: Text to replace with
    
    Returns:
        Result with replacement count
    """
    try:
        file_path = get_file_path(filename)
        
        if not file_path.exists():
            return {"success": False, "error": f"File not found: {filename}"}
        
        doc = Document(str(file_path))
        count = 0
        
        for para in doc.paragraphs:
            for run in para.runs:
                if search_text in run.text:
                    run.text = run.text.replace(search_text, replace_text)
                    count += run.text.count(replace_text)
        
        doc.save(str(file_path))
        
        return {
            "success": True,
            "message": f"Replaced {count} occurrences",
            "replacement_count": count
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


# ==================== Resources ====================

@mcp.resource("file://documents")
def list_documents_resource() -> str:
    """Get list of all Word documents."""
    result = list_documents()
    return json.dumps(result, indent=2)


# ==================== Prompts ====================

@mcp.prompt()
def help_prompt() -> list[dict]:
    """Get help for Word document operations."""
    return [{
        "role": "user",
        "content": """Help me with Word document operations:
1. How to create a document
2. How to add tables
3. How to format text
4. How to search and replace"""
    }]


# Run server
if __name__ == "__main__":
    mcp.run()

