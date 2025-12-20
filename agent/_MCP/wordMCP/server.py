"""
Word MCP Server with SSE Transport

启动方式:
    python server.py

服务器会在 http://localhost:8080 启动 SSE 端点
"""

import asyncio
import json
import logging
from typing import Optional, List
from pathlib import Path
from datetime import datetime
import re

from fastapi import FastAPI, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# FastAPI 应用
app = FastAPI(title="Word MCP Server")

# CORS 配置
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 文档目录
WORD_DIR = Path("word")
WORD_DIR.mkdir(exist_ok=True)


# ==================== 工具函数 ====================

def get_file_path(filename: str) -> Path:
    """获取文件完整路径"""
    path = Path(filename)
    if not path.is_absolute():
        path = WORD_DIR / filename
    if not str(path).endswith('.docx'):
        path = Path(str(path) + '.docx')
    return path


# ==================== MCP 工具实现 ====================

TOOLS = {
    "create_document": {
        "description": "创建新的 Word 文档",
        "parameters": {
            "filename": {"type": "string", "description": "文件名（可选）"},
            "title": {"type": "string", "description": "文档标题（可选）"},
            "content": {"type": "string", "description": "初始内容（可选）"}
        }
    },
    "read_document": {
        "description": "读取 Word 文档内容",
        "parameters": {
            "filename": {"type": "string", "description": "文件名", "required": True}
        }
    },
    "update_document": {
        "description": "更新 Word 文档",
        "parameters": {
            "filename": {"type": "string", "required": True},
            "action": {"type": "string", "enum": ["append", "insert", "replace", "add_heading"], "required": True},
            "content": {"type": "string"},
            "paragraph_index": {"type": "integer"}
        }
    },
    "delete_document": {
        "description": "删除 Word 文档",
        "parameters": {
            "filename": {"type": "string", "required": True}
        }
    },
    "list_documents": {
        "description": "列出所有 Word 文档",
        "parameters": {}
    },
    "add_table": {
        "description": "向文档添加表格",
        "parameters": {
            "filename": {"type": "string", "required": True},
            "table_data": {"type": "array", "description": "表格数据（二维数组）", "required": True},
            "title": {"type": "string"}
        }
    },
    "search_replace": {
        "description": "搜索并替换文本",
        "parameters": {
            "filename": {"type": "string", "required": True},
            "search_text": {"type": "string", "required": True},
            "replace_text": {"type": "string", "required": True}
        }
    }
}


def create_document(filename: str = None, title: str = None, content: str = None) -> dict:
    """创建新文档"""
    try:
        if not filename:
            filename = f"document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        file_path = get_file_path(filename)
        doc = Document()
        
        if title:
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if content:
            for line in content.split('\n'):
                doc.add_paragraph(line.strip() if line.strip() else "")
        
        doc.save(str(file_path))
        
        return {
            "success": True,
            "message": "文档创建成功",
            "file_path": str(file_path),
            "file_size": file_path.stat().st_size
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


def read_document(filename: str) -> dict:
    """读取文档"""
    try:
        file_path = get_file_path(filename)
        
        if not file_path.exists():
            return {"success": False, "error": f"文件不存在: {filename}"}
        
        doc = Document(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs]
        
        tables = []
        for table in doc.tables:
            table_data = [[cell.text for cell in row.cells] for row in table.rows]
            tables.append(table_data)
        
        return {
            "success": True,
            "file_path": str(file_path),
            "paragraphs": paragraphs,
            "paragraph_count": len(paragraphs),
            "tables": tables,
            "table_count": len(tables),
            "full_text": "\n".join(paragraphs)
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


def update_document(filename: str, action: str, content: str = None, paragraph_index: int = None) -> dict:
    """更新文档"""
    try:
        file_path = get_file_path(filename)
        
        if not file_path.exists():
            return {"success": False, "error": f"文件不存在: {filename}"}
        
        doc = Document(str(file_path))
        
        if action == "append" and content:
            for line in content.split('\n'):
                doc.add_paragraph(line)
        elif action == "add_heading" and content:
            doc.add_heading(content, level=2)
        elif action == "insert" and content and paragraph_index is not None:
            if paragraph_index < len(doc.paragraphs):
                doc.paragraphs[paragraph_index].insert_paragraph_before(content)
        elif action == "replace" and paragraph_index is not None:
            if paragraph_index < len(doc.paragraphs):
                para = doc.paragraphs[paragraph_index]
                para.clear()
                para.add_run(content or "")
        else:
            return {"success": False, "error": "无效的操作或缺少参数"}
        
        doc.save(str(file_path))
        return {"success": True, "message": "文档更新成功", "action": action}
    except Exception as e:
        return {"success": False, "error": str(e)}


def delete_document(filename: str) -> dict:
    """删除文档"""
    try:
        file_path = get_file_path(filename)
        
        if not file_path.exists():
            return {"success": False, "error": f"文件不存在: {filename}"}
        
        file_path.unlink()
        return {"success": True, "message": "文档删除成功"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def list_documents() -> dict:
    """列出所有文档"""
    try:
        docs = []
        for file in WORD_DIR.glob("*.docx"):
            stat = file.stat()
            docs.append({
                "name": file.name,
                "path": str(file),
                "size": stat.st_size,
                "modified": datetime.fromtimestamp(stat.st_mtime).isoformat()
            })
        return {"success": True, "count": len(docs), "documents": docs}
    except Exception as e:
        return {"success": False, "error": str(e)}


def add_table(filename: str, table_data: list, title: str = None) -> dict:
    """添加表格"""
    try:
        file_path = get_file_path(filename)
        
        if not file_path.exists():
            return {"success": False, "error": f"文件不存在: {filename}"}
        
        doc = Document(str(file_path))
        
        if title:
            doc.add_heading(title, level=2)
        
        if table_data:
            rows = len(table_data)
            cols = max(len(row) for row in table_data)
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Light Grid Accent 1'
            
            for i, row_data in enumerate(table_data):
                for j, cell_data in enumerate(row_data):
                    if j < cols:
                        table.rows[i].cells[j].text = str(cell_data)
        
        doc.save(str(file_path))
        return {"success": True, "message": "表格添加成功"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def search_replace(filename: str, search_text: str, replace_text: str) -> dict:
    """搜索替换"""
    try:
        file_path = get_file_path(filename)
        
        if not file_path.exists():
            return {"success": False, "error": f"文件不存在: {filename}"}
        
        doc = Document(str(file_path))
        count = 0
        
        for para in doc.paragraphs:
            for run in para.runs:
                if search_text in run.text:
                    run.text = run.text.replace(search_text, replace_text)
                    count += 1
        
        doc.save(str(file_path))
        return {"success": True, "message": f"替换了 {count} 处", "count": count}
    except Exception as e:
        return {"success": False, "error": str(e)}


# 工具映射
TOOL_HANDLERS = {
    "create_document": create_document,
    "read_document": read_document,
    "update_document": update_document,
    "delete_document": delete_document,
    "list_documents": list_documents,
    "add_table": add_table,
    "search_replace": search_replace,
}


# ==================== API 端点 ====================

class ToolCallRequest(BaseModel):
    """工具调用请求"""
    tool: str
    params: dict = {}

class AgentRequest(BaseModel):
    """编排请求：一次 SSE 流里执行多步工具调用"""
    query: str
    title: Optional[str] = None
    filename: Optional[str] = None

def _sanitize_filename(name: str) -> str:
    # 保留中英文、数字、下划线、短横线；空白转下划线
    s = re.sub(r"\s+", "_", name.strip())
    s = re.sub(r"[^\w\u4e00-\u9fa5_-]+", "", s)
    return s[:40] if s else ""

def _react_doc_plan(title: str, filename: str) -> list[dict]:
    """
    生成一个“可执行的工具计划”：每一步是 {tool, params, label}
    说明：这里不接 LLM，直接用模板生成内容，确保可用、可重复。
    """
    intro = (
        f"《{title}》\n"
        f"\n"
        f"本文档用于快速了解 {title} 的定位、核心概念与实践建议，并给出一个最小示例。\n"
    )

    sections: list[tuple[str, str]] = [
        ("一、React 是什么", "\n".join([
            "React 是一个用于构建用户界面（UI）的 JavaScript 库，核心思想是用组件化的方式组织界面。",
            "它强调声明式编程：你描述 UI 在某个状态下应该长什么样，框架负责在状态变化时高效更新界面。",
            "React 本身更关注“视图层”，通常会配合路由、状态管理、构建工具等组成完整工程方案。"
        ])),
        ("二、核心概念速览", "\n".join([
            "1) 组件（Component）：将 UI 拆分为可复用单元。",
            "2) JSX：一种 JavaScript 语法扩展，用于更直观地描述 UI 结构。",
            "3) Props / State：props 用于父传子数据；state 用于组件内部状态。",
            "4) Hooks：函数组件中复用状态逻辑（如 useState/useEffect/useMemo）。",
            "5) 渲染与更新：React 会在状态变化后计算新的 UI，并将差异更新到真实 DOM。"
        ])),
        ("三、Hooks 使用建议", "\n".join([
            "useState：管理局部状态；状态更新尽量保持不可变数据结构。",
            "useEffect：处理副作用（数据请求、订阅、DOM 操作）；注意清理函数避免泄漏。",
            "useMemo / useCallback：用于性能优化，避免不必要的重新计算与函数重建（不要过度使用）。",
            "自定义 Hook：把可复用逻辑封装成 useXxx，提升可维护性。"
        ])),
        ("四、生态与工程化", "\n".join([
            "路由：React Router（Web 应用常用）。",
            "状态管理：Redux Toolkit / Zustand / Jotai 等（按复杂度选择）。",
            "数据请求与缓存：TanStack Query（React Query）。",
            "构建工具：Vite / Webpack；框架化方案可选 Next.js（SSR/SSG）。",
            "类型系统：TypeScript 与 React 搭配能显著提升大型项目稳定性。"
        ])),
        ("五、一个最小示例（函数组件 + Hooks）", "\n".join([
            "下面示例展示一个计数器：",
            "",
            "```tsx",
            "import { useState } from 'react';",
            "",
            "export function Counter() {",
            "  const [count, setCount] = useState(0);",
            "  return (",
            "    <div>",
            "      <p>count: {count}</p>",
            "      <button onClick={() => setCount((c) => c + 1)}>+1</button>",
            "    </div>",
            "  );",
            "}",
            "```"
        ])),
        ("六、最佳实践清单", "\n".join([
            "组件职责单一：展示组件与容器组件适度分离。",
            "数据流清晰：优先自上而下传递，避免过早引入全局状态。",
            "副作用可控：useEffect 依赖数组准确；异步请求注意取消/忽略过期响应。",
            "性能优化有依据：先用性能分析工具定位，再用 memo/useMemo/useCallback 优化。",
            "可测试性：复杂逻辑提取为纯函数/自定义 Hook，便于单测。"
        ])),
        ("七、总结", "\n".join([
            f"{title} 适合构建组件化、可维护的大型前端应用。",
            "建议从函数组件 + Hooks 入手，配合 TypeScript 与现代构建工具（如 Vite）形成稳定工程实践。"
        ])),
    ]

    table = [
        ["维度", "React", "Vue", "Angular"],
        ["定位", "UI 库", "渐进式框架", "全家桶框架"],
        ["学习曲线", "中等（JS/TS + Hooks）", "较平滑", "较陡（概念多）"],
        ["生态", "非常丰富", "丰富", "偏企业级"],
    ]

    plan: list[dict] = [
        {
            "tool": "create_document",
            "label": "创建文档",
            "params": {"filename": filename, "title": title, "content": intro},
        }
    ]

    for heading, body in sections:
        plan.append({"tool": "update_document", "label": f"添加章节：{heading}", "params": {
            "filename": filename, "action": "add_heading", "content": heading
        }})
        plan.append({"tool": "update_document", "label": f"写入内容：{heading}", "params": {
            "filename": filename, "action": "append", "content": body
        }})

    plan.append({"tool": "add_table", "label": "插入对比表格", "params": {
        "filename": filename, "title": "React / Vue / Angular 对比（简表）", "table_data": table
    }})

    return plan


@app.get("/")
async def root():
    """服务器状态"""
    return {
        "name": "Word MCP Server",
        "version": "1.0.0",
        "status": "running",
        "endpoints": {
            "tools": "/tools",
            "call": "/call (POST)",
            "sse": "/sse",
            "documents": "/documents"
        }
    }


@app.get("/tools")
async def get_tools():
    """获取可用工具列表"""
    return {"tools": TOOLS}


@app.post("/call")
async def call_tool(request: ToolCallRequest):
    """调用工具"""
    tool_name = request.tool
    params = request.params
    
    if tool_name not in TOOL_HANDLERS:
        return {"success": False, "error": f"未知工具: {tool_name}"}
    
    handler = TOOL_HANDLERS[tool_name]
    result = handler(**params)
    
    logger.info(f"工具调用: {tool_name}({params}) -> {result.get('success')}")
    return result


@app.get("/documents")
async def get_documents():
    """获取文档列表"""
    return list_documents()


@app.get("/sse")
async def sse_endpoint(request: Request):
    """SSE 端点 - 用于实时事件流"""
    
    async def event_generator():
        # 发送连接确认
        yield f"data: {json.dumps({'type': 'connected', 'message': 'SSE 连接成功'})}\n\n"
        
        # 发送可用工具列表
        yield f"data: {json.dumps({'type': 'tools', 'tools': list(TOOLS.keys())})}\n\n"
        
        # 保持连接
        while True:
            if await request.is_disconnected():
                break
            
            # 发送心跳
            yield f"data: {json.dumps({'type': 'heartbeat', 'time': datetime.now().isoformat()})}\n\n"
            await asyncio.sleep(30)
    
    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no"
        }
    )


@app.post("/sse/call")
async def sse_call_tool(request: ToolCallRequest):
    """
    SSE 方式调用工具
    返回 SSE 流式响应
    """
    
    async def stream_response():
        tool_name = request.tool
        params = request.params
        
        # 发送开始事件
        yield f"data: {json.dumps({'type': 'start', 'tool': tool_name})}\n\n"
        
        if tool_name not in TOOL_HANDLERS:
            yield f"data: {json.dumps({'type': 'error', 'error': f'未知工具: {tool_name}'})}\n\n"
            return
        
        # 执行工具
        handler = TOOL_HANDLERS[tool_name]
        result = handler(**params)
        
        # 发送结果
        yield f"data: {json.dumps({'type': 'result', 'data': result})}\n\n"
        
        # 发送完成事件
        yield f"data: {json.dumps({'type': 'done'})}\n\n"
    
    return StreamingResponse(
        stream_response(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive"
        }
    )

@app.post("/sse/agent")
async def sse_agent(request: AgentRequest, http_request: Request):
    """
    真·多步 SSE：在同一条 SSE 响应里，按计划连续调用多个工具并持续推送进度。
    """

    async def stream_response():
        # 解析标题/文件名
        title = (request.title or "").strip()
        if not title:
            # 极简兜底：从 query 里提取（不做复杂 NLP）
            title = request.query.strip() or "新文档"

        base_name = _sanitize_filename(request.filename or title)
        if not base_name or len(base_name) < 3:
            base_name = f"doc_{int(datetime.now().timestamp())}"
        filename = f"{base_name}_{int(datetime.now().timestamp() * 1000)}"

        # 发送编排开始事件
        yield f"data: {json.dumps({'type': 'start', 'tool': 'agent', 'message': '开始多步生成文档'})}\n\n"
        await asyncio.sleep(0)

        # 这里先对 React 做一个高质量模板；其他主题也能复用（只是内容更通用）
        plan = _react_doc_plan(title=title, filename=filename)

        for idx, step in enumerate(plan, start=1):
            if await http_request.is_disconnected():
                break

            tool_name = step["tool"]
            params = step["params"]
            label = step.get("label") or tool_name

            yield f"data: {json.dumps({'type': 'progress', 'message': f'[{idx}/{len(plan)}] {label}'})}\n\n"
            yield f"data: {json.dumps({'type': 'start', 'tool': tool_name, 'step': idx, 'label': label})}\n\n"

            if tool_name not in TOOL_HANDLERS:
                yield f"data: {json.dumps({'type': 'error', 'error': f'未知工具: {tool_name}', 'step': idx})}\n\n"
                return

            handler = TOOL_HANDLERS[tool_name]
            result = handler(**params)
            yield f"data: {json.dumps({'type': 'result', 'data': result, 'tool': tool_name, 'step': idx, 'label': label})}\n\n"
            await asyncio.sleep(0)

        yield f"data: {json.dumps({'type': 'done', 'data': {'filename': filename, 'title': title}})}\n\n"

    return StreamingResponse(
        stream_response(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


if __name__ == "__main__":
    import uvicorn
    
    print("=" * 50)
    print("Word MCP Server (SSE)")
    print("=" * 50)
    print(f"文档目录: {WORD_DIR.absolute()}")
    print(f"服务地址: http://localhost:8080")
    print("=" * 50)
    print("\n可用端点:")
    print("  GET  /          - 服务器状态")
    print("  GET  /tools     - 工具列表")
    print("  POST /call      - 调用工具")
    print("  GET  /sse       - SSE 连接")
    print("  POST /sse/call  - SSE 调用工具")
    print("  POST /sse/agent - SSE 多步编排（真·流式）")
    print("  GET  /documents - 文档列表")
    print("=" * 50)
    
    uvicorn.run(app, host="0.0.0.0", port=8080)

